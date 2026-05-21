# -*- coding: utf-8 -*-
"""将 journal/5.9版本介绍.md 导出为 Word .docx（与 md 同目录）。"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt


def repo_root() -> Path:
    return Path(__file__).resolve().parents[1]


def parse_inline_bold(paragraph, text: str) -> None:
    """段落内粗体 **xxx**."""
    parts = re.split(r"(\*\*.+?\*\*)", text)
    run0 = True
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(11)
        elif part:
            run = paragraph.add_run(part)
            run.font.size = Pt(11)
        run0 = False


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text.strip(), level=level)
    for r in p.runs:
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
            "Microsoft YaHei",
        )


def is_table_sep(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|"):
        return False
    inner = s.strip("|")
    return bool(re.fullmatch(r"[\s\-:|]+", inner))


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines):
        ln = lines[i].strip()
        if not ln.startswith("|"):
            break
        if is_table_sep(ln):
            i += 1
            continue
        cells = [c.strip() for c in ln.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def flush_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < cols:
            r.append("")
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            parse_inline_bold(table.rows[ri].cells[ci].paragraphs[0], cell_text)


def convert(md_path: Path, out_path: Path) -> None:
    raw = md_path.read_text(encoding="utf-8")
    lines = raw.splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()

        if not s:
            i += 1
            continue
        if s == "---":
            i += 1
            continue

        if s.startswith("|"):
            rows, ni = parse_table(lines, i)
            flush_table(doc, rows)
            i = ni
            doc.add_paragraph()
            continue

        if s.startswith("### "):
            add_heading(doc, s[4:], level=3)
            i += 1
            continue
        if s.startswith("## "):
            add_heading(doc, s[3:], level=2)
            i += 1
            continue
        if s.startswith("# "):
            add_heading(doc, s[2:], level=1)
            i += 1
            continue

        # 有序 / 无序列表
        bullet = False
        body = s
        if body.startswith("- "):
            bullet = True
            body = body[2:]
        elif re.match(r"^\d+\.\s+", body):
            body = re.sub(r"^\d+\.\s+", "", body)

        p = doc.add_paragraph(style="List Bullet" if bullet else "Normal")
        parse_inline_bold(p, body)
        i += 1

    doc.save(out_path)


def main() -> None:
    root = repo_root()
    md = root / "src/main/resources/com/example/monopoly_deal_game/journal/5.9版本介绍.md"
    docx = root / "src/main/resources/com/example/monopoly_deal_game/journal/5.9版本介绍.docx"
    if not md.is_file():
        raise SystemExit(f"找不到 Markdown: {md}")
    convert(md, docx)
    print(f"已生成: {docx}")


if __name__ == "__main__":
    main()
